"""
A unified script to process raw and processed data for fine-tuning.

This script extracts and transforms data from various sources:
- .docx files for monolingual text and grammar instructions.
- .jsonl files for translations, monolingual text, and grammar instructions.

It formats all data into a consistent JSONL format with 'instruction',
'completion', and 'task_type' fields, ready for fine-tuning a language model.
"""

import json
import re
from pathlib import Path

from docx import Document

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
DOCX_GRAMMAR_PATH = Path("data/raw/grammar/RUNYAKITARA LANGUAGE STUDIES.docx")
MONOLINGUAL_DOCX_DIR = Path("data/raw/monolingual")
LOOKUP_JSONL_PATH = Path("data/processed/lookup_data.jsonl")
MONOLINGUAL_JSONL_PATH = Path("data/processed/monolingual_data.jsonl")
INSTRUCTION_JSONL_PATH = Path("data/processed/instruction_data.jsonl")

OUTPUT_DIR = Path("data/processed")
OUTPUT_FILE = OUTPUT_DIR / "final_dataset.jsonl"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------

def _is_bold(para):
    """Check if the entire paragraph is bold."""
    return all(run.bold for run in para.runs)


def _is_heading(para):
    """
    Determines if a paragraph is a heading.
    A heading is defined as a line that is either bold or all uppercase.
    """
    text = para.text.strip()
    if not text:
        return False
    return _is_bold(para) or text.isupper()


def _chunk_text(text, min_words=20, max_words=256):
    """Splits text into chunks (paragraphs) and then into smaller pieces."""
    chunks = []
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    for para in paragraphs:
        words = para.split()
        if len(words) < min_words:
            continue

        # If paragraph is too long, split it.
        if len(words) > max_words:
            midpoint = len(words) // 2
            part1 = " ".join(words[:midpoint])
            part2 = " ".join(words[midpoint:])
            chunks.append((part1, part2))
        else:
            # If paragraph is within limits, split into prompt/completion
            midpoint = len(words) // 2
            if midpoint == 0:
                continue
            prompt = " ".join(words[:midpoint])
            completion = " ".join(words[midpoint:])
            chunks.append((prompt, completion))
    return chunks


# ---------------------------------------------------------------------------
# Data Extraction from JSONL files
# ---------------------------------------------------------------------------

def process_lookup_data(filepath):
    """Processes lookup data for translation tasks."""
    data = []
    print(f"Processing translation data from {filepath}...")
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            entry = json.loads(line)
            # Basic cleaning: remove slashes and extra context.
            english_term = entry.get("english_term", "").strip()
            runyoro_term = entry.get("runyoro_term", "").strip()

            if not english_term or not runyoro_term:
                continue

            # A simple heuristic to check if the term is a sentence-like fragment
            if len(english_term.split()) > 2 and not english_term.startswith('/'):
                 data.append({
                    "instruction": f"Translate the following English sentence to Rutooro: {english_term}",
                    "completion": runyoro_term,
                    "task_type": "translation"
                })
    print(f"Extracted {len(data)} translation entries.")
    return data

def process_monolingual_jsonl(filepath):
    """Processes monolingual data for generation tasks."""
    data = []
    print(f"Processing monolingual data from {filepath}...")
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            entry = json.loads(line)
            text = entry.get("text", "")
            for prompt, completion in _chunk_text(text):
                data.append({
                    "instruction": f"Continue the following story in Rutooro: {prompt}",
                    "completion": completion,
                    "task_type": "monolingual_generation"
                })
    print(f"Extracted {len(data)} monolingual generation entries.")
    return data

def process_instruction_jsonl(filepath):
    """Processes existing instruction data."""
    data = []
    print(f"Processing grammar instruction data from {filepath}...")
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            entry = json.loads(line)
            instruction = entry.get("instruction")
            completion = entry.get("completion")
            if instruction and completion:
                data.append({
                    "instruction": instruction,
                    "completion": completion,
                    "task_type": "grammar_instruction"
                })
    print(f"Extracted {len(data)} grammar instruction entries.")
    return data


# ---------------------------------------------------------------------------
# Data Extraction from DOCX files
# ---------------------------------------------------------------------------

def extract_grammar_from_docx(paragraphs, used_paras):
    """Extracts grammar instructions and examples from the document."""
    instruction_data = []
    i = 0
    while i < len(paragraphs):
        if i in used_paras:
            i += 1
            continue

        para = paragraphs[i]
        if _is_heading(para):
            instruction = para.text.strip()
            instruction_paras = {i}

            completion_lines = []
            completion_paras = set()

            j = i + 1
            while j < len(paragraphs):
                if j in used_paras or _is_heading(paragraphs[j]):
                    break
                comp_para = paragraphs[j]
                comp_text = comp_para.text.strip()
                if not comp_text:
                    break
                completion_lines.append(comp_text)
                completion_paras.add(j)
                j += 1

            if completion_lines:
                completion = " ".join(completion_lines)
                instruction_data.append({
                    "instruction": instruction,
                    "completion": completion,
                    "task_type": "grammar_instruction"
                })
                used_paras.update(instruction_paras)
                used_paras.update(completion_paras)
            i = j
        else:
            i += 1
    print(f"Extracted {len(instruction_data)} grammar entries from DOCX.")
    return instruction_data


def extract_monolingual_from_docx(paragraphs=None, used_paras=None, folder_path=None):
    """Extracts coherent paragraphs of text for monolingual generation."""
    monolingual_data = []

    text_sources = []
    if paragraphs and used_paras is not None:
        # From a list of paragraphs, skipping used ones
        for i, para in enumerate(paragraphs):
            if i in used_paras:
                continue
            text_sources.append(para.text.strip())

    if folder_path:
        # From a folder of .docx files
        for doc_path in folder_path.glob("*.docx"):
            document = Document(doc_path)
            for para in document.paragraphs:
                text_sources.append(para.text.strip())

    for text in text_sources:
        for prompt, completion in _chunk_text(text):
            monolingual_data.append({
                "instruction": f"Continue the following story in Rutooro: {prompt}",
                "completion": completion,
                "task_type": "monolingual_generation"
            })

    print(f"Extracted {len(monolingual_data)} monolingual generation entries from DOCX sources.")
    return monolingual_data


# ---------------------------------------------------------------------------
# Main Orchestration
# ---------------------------------------------------------------------------

def main():
    """Main function to orchestrate the data extraction and processing."""
    all_data = []

    # 1. Process JSONL files
    all_data.extend(process_lookup_data(LOOKUP_JSONL_PATH))
    all_data.extend(process_monolingual_jsonl(MONOLINGUAL_JSONL_PATH))
    all_data.extend(process_instruction_jsonl(INSTRUCTION_JSONL_PATH))

    # 2. Process DOCX files
    print(f"Loading grammar document: {DOCX_GRAMMAR_PATH}")
    if DOCX_GRAMMAR_PATH.exists():
        document = Document(DOCX_GRAMMAR_PATH)
        paragraphs = document.paragraphs
        used_paras = set()

        all_data.extend(extract_grammar_from_docx(paragraphs, used_paras))
        all_data.extend(extract_monolingual_from_docx(paragraphs=paragraphs, used_paras=used_paras))
    else:
        print(f"Warning: Grammar DOCX not found at {DOCX_GRAMMAR_PATH}")

    # 3. Extract monolingual text from dedicated folder
    if MONOLINGUAL_DOCX_DIR.exists():
        all_data.extend(extract_monolingual_from_docx(folder_path=MONOLINGUAL_DOCX_DIR))
    else:
        print(f"Warning: Monolingual DOCX directory not found at {MONOLINGUAL_DOCX_DIR}")


    # 4. Write final dataset
    print(f"Writing final dataset with {len(all_data)} entries to {OUTPUT_FILE}")
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        for item in all_data:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")

    print("Processing complete.")


if __name__ == "__main__":
    main()
